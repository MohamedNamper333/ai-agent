"""PDF, DOCX, Excel, and Image analysis tools"""

import base64
import io
import os
import re
from pathlib import Path
from typing import Optional


class DocumentTools:
    @staticmethod
    def read_pdf(file_path: str, max_pages: int = 20) -> str:
        """Extract and return text from each page of a PDF file."""
        p = Path(file_path)
        if not p.exists():
            return f"Error: File not found: {file_path}"

        try:
            import PyPDF2
            with open(p, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                total = len(reader.pages)
                pages_to_read = min(total, max_pages)
                parts = []
                for i in range(pages_to_read):
                    text = reader.pages[i].extract_text()
                    text = re.sub(r'\s+', ' ', text).strip()
                    parts.append(f"--- Page {i+1}/{total} ---\n{text}")

                result = "\n\n".join(parts)
                if max_pages < total:
                    result += f"\n\n... ({total - max_pages} more pages truncated)"

                return f"PDF: {p.name} ({total} pages)\n\n{result[:10000]}"
        except ImportError:
            return "Error: Install PyPDF2 (pip install PyPDF2)"
        except Exception as e:
            return f"Error reading PDF: {e}"

    @staticmethod
    def read_docx(file_path: str) -> str:
        """Extract paragraphs and tables from a Word document."""
        p = Path(file_path)
        if not p.exists():
            return f"Error: File not found: {file_path}"

        try:
            import docx
            doc = docx.Document(str(p))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n".join(paragraphs)

            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                tables.append("\n".join(rows))

            result = f"DOCX: {p.name}\n\nParagraphs:\n{text[:5000]}"
            if tables:
                result += f"\n\nTables ({len(tables)}):\n"
                result += "\n\n".join(tables[:5])
            return result[:10000]
        except ImportError:
            return "Error: Install python-docx (pip install python-docx)"
        except Exception as e:
            return f"Error reading DOCX: {e}"

    @staticmethod
    def analyze_image(file_path: str) -> str:
        """Return image metadata and optional OCR text."""
        p = Path(file_path)
        if not p.exists():
            return f"Error: File not found: {file_path}"

        try:
            from PIL import Image
            img = Image.open(p)
            info = (
                f"Image: {p.name}\n"
                f"Format: {img.format}\n"
                f"Size: {img.size[0]}x{img.size[1]} pixels\n"
                f"Mode: {img.mode}\n"
                f"File size: {p.stat().st_size:,} bytes\n"
            )

            try:
                import pytesseract
                text = pytesseract.image_to_string(img)
                text = text.strip()
                if text:
                    info += f"\nOCR Text:\n{text[:2000]}\n"
            except ImportError:
                info += "\n(OCR: install pytesseract for text extraction)"

            return info
        except ImportError:
            return "Error: Install Pillow (pip install Pillow)"
        except Exception as e:
            return f"Error analyzing image: {e}"

    @staticmethod
    def ocr_image(file_path: str) -> str:
        """Extract and return text from an image using OCR."""
        p = Path(file_path)
        if not p.exists():
            return f"Error: File not found: {file_path}"

        try:
            from PIL import Image
            import pytesseract
            img = Image.open(p)
            text = pytesseract.image_to_string(img)
            text = text.strip()
            if text:
                return f"OCR Text from {p.name}:\n{text[:5000]}"
            return f"No text found in {p.name}"
        except ImportError:
            return "Error: Install Pillow + pytesseract"
        except Exception as e:
            return f"Error in OCR: {e}"

    @staticmethod
    def read_excel(file_path: str, sheet: str = "") -> str:
        """Read an Excel file and return the first sheet contents."""
        p = Path(file_path)
        if not p.exists():
            return f"Error: File not found: {file_path}"

        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)

            sheet_names = wb.sheetnames
            target_sheet = sheet if sheet and sheet in sheet_names else sheet_names[0]
            ws = wb[target_sheet]

            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                return f"Excel file is empty: {file_path}"

            result = [
                f"Excel: {p.name}",
                f"Sheet: {target_sheet}",
                f"Rows: {len(rows)-1}",
                "",
                "First 5 rows:",
            ]
            for row in rows[:5]:
                result.append(f"  {list(row)}")

            wb.close()
            return "\n".join(result)

        except ImportError:
            return "Error: Install openpyxl (pip install openpyxl)"
        except Exception as e:
            return f"Error reading Excel: {e}"

    @staticmethod
    def html_to_text(html_content: str) -> str:
        """Strip HTML tags and return clean plain text."""
        try:
            import re
            text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()

            if len(text) > 10000:
                text = text[:10000] + "\n... (truncated)"
            return text
        except Exception as e:
            return f"Error: {e}"
